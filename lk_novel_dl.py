import time

import cv2
import imageio
import numpy as np
from bs4 import BeautifulSoup
import requests
import os
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import pickle
import shutil
from io import BytesIO
from docx import shared
from argparse import ArgumentParser


description = f'''
LK Novel Download Tool
LK小说下载工具

示例
python3 {os.path.basename(__file__)} https://www.lightnovel.us/cn/detail/495565

下载完成后，输出文件在当前目录下，默认名字为 out.docx

-------------------------------
'''


print(description)

parser = ArgumentParser(description='')

parser.add_argument('url', type=str, help='这是填入轻国小说的url')
parser.add_argument('--no-title', action='store_true', default=False, help='设定该标志可关闭自动的标题检测和标题生成')
parser.add_argument('--no-cache', action='store_true', default=False, help='设定该标志可以关闭缓存，将会始终使用在线链接')
parser.add_argument('--replace-txt', type=str, default=None, help='高级用法，替换url指向资源，使用方法参见 --help-replace-txt')
parser.add_argument('--help-replace-txt', action='store_true', default=False, help='显示 --replace-txt 怎么用')
parser.add_argument('--out', type=str, default='out.docx', help='输出文件的路径')

parser = parser.parse_args()

# url = 'https://www.lightnovel.us/cn/detail/597883'
# url = 'https://www.lightnovel.us/cn/detail/495565'
# url = 'https://www.lightnovel.us/detail/802366'
# url = 'https://www.lightnovel.us/cn/detail/594945'
# url = 'https://www.lightnovel.us/cn/detail/685846'
# url = 'https://www.lightnovel.us/cn/detail/509587'


url = parser.url
no_cache = parser.no_cache
no_title = parser.no_title
replace_txt = parser.replace_txt
out_path = parser.out


if parser.help_replace_txt:
    help_txt = '''
    本功能用来替换指定url为本地资源。
    某些轻国小说页面结构有问题，例如 https://www.lightnovel.us/cn/detail/336435
    需要把网页下载下来，然后手动修正html结构，不然会缺失某些项。
    replace-txt 是一个普通的文本文件，可以用notepad编辑和生成
    文件结构为 行式结构，行间使用回车符分隔，行内使用空格符分隔。第一个列为请求的url，第二列为要替换的本地资源的路径
    
    replace-txt 结构
    
    请求的url1 本地资源路径1
    请求的url2 本地资源路径2
    '''
    print(help_txt)
    exit(0)


# url资源替换字典
replace_dict = {}


# 这里解析 replace_txt
if replace_txt is not None:
    print('启用了替换url为本地资源功能')
    lines = open(replace_txt, 'r').read().splitlines(False)
    for li in lines:
        its = li.split(' ', 1)
        if len(its) != 2:
            print(f'Warning! Found bad line in replace_txt. Will ignore this line. {li}')
            continue
        replace_dict[its[0]] = its[1]
    print('当前有效replace_txt设置')
    print(replace_dict)
    print()


ua = 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/86.0.4240.198 Safari/537.36'

# 代理设置
proxies = {
    # 'http': 'http://192.168.0.76:1080',
    # 'https': 'https://192.168.0.76:1080'
}

cache_file = 'cache.pkl'

if not os.path.isfile(cache_file):
    pickle.dump({}, open(cache_file, 'wb'))

cache_dict: dict = pickle.load(open(cache_file, 'rb'))

session = requests.Session()


def get_url(url: str, ref=None):
    '''
    下载函数
    :param url:
    :param ref:
    :return:
    '''
    # url资源替换处
    if url in replace_dict:
        d_path = replace_dict[url]
        if os.path.isfile(d_path):
            return open(d_path, 'rb').read()
        else:
            print(f'Warning! 指定本地资源不存在或不可读！本次替换将被忽略。 {d_path}')

    # 一些简单的处理，可以去掉这个图站的水印
    if url.find('i1137.photobucket.com') != -1:
        ref = url.replace('i1137.photobucket.com', 'hosting.photobucket.com')
    elif url.find('i1138.photobucket.com') != -1:
        ref = url.replace('i1138.photobucket.com', 'hosting.photobucket.com')

    assert url.startswith('http'), f'Error! Bad url! {url}'
    if url not in cache_dict or no_cache:
        success = False
        for _ in range(3):
            # 必须使用等待，如果请求速度过快，会被ban。
            time.sleep(1)
            try:
                if ref is None:
                    q = session.get(url, headers={'user-agent': ua}, timeout=5, proxies=proxies)
                else:
                    q = session.get(url, headers={'referer': ref, 'user-agent': ua}, timeout=5, proxies=proxies)
                if q.status_code == 200:
                    success = True
                    break
                print(f'Bad {q.status_code} url {url}')
            except Exception as e:
                print(e, 'url', url)

        if not success:
            return None
        else:
            print('url success')

        cache_dict[url] = q.content
        shutil.move(cache_file, cache_file + '.bak')
        pickle.dump(cache_dict, open(cache_file, 'wb'))

    return cache_dict[url]


text = get_url(url).decode('utf8')

w = BeautifulSoup(text, features='lxml')

content = w.find(id='article-main-contents')

document = Document()

# 最后一段，全局变量
last_p = None
# 连续空行计数
find_br_count = 0


def new_pa(doc, title_level=None):
    '''
    新增一个段
    :return:
    '''
    if title_level is not None:
        # level 0 不会被识别为标题
        # 要使用level 1 和以上的，才能识别为章节标题
        p = doc.add_heading('', level=title_level)  # 这里不填标题内容
    else:
        p = doc.add_paragraph()
    paragraph_format = p.paragraph_format
    paragraph_format.line_spacing = 1.1
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    return p


def new_run(p, text):
    run = p.add_run(text)
    run.font.name = '等线'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '等线')
    run.font.color.rgb = RGBColor(0, 0, 0)
    return run


def add_title(doc, t):
    '''
    增加一个新标题，设定为章节头
    :param doc:
    :param t:
    :return:
    '''
    global last_p
    # 要使用level 1的，才能识别为章节标题
    if no_title:
        p = new_pa(doc)
    else:
        p = new_pa(doc, 1)
    run = new_run(p, t)
    last_p = p


def add_failure_url(doc, t):
    '''
    增加一个失败链接文本，风格为2级标题，颜色为红色，用来快速找到失败链接和手动更新
    :param doc:
    :param t:
    :return:
    '''
    global last_p
    p = new_pa(doc, 2)
    run = new_run(p, t)
    run.font.color.rgb = RGBColor(255, 0, 0)
    last_p = p


def add_text(doc, t):
    '''
    增加一个新段，并填充文字
    :param doc:
    :param t:
    :return:
    '''
    global last_p
    p = new_pa(doc)
    run = new_run(p, t)
    last_p = p


def add_last_text(doc, t):
    '''
    在最后段的后面，继续添加文字
    如果最后段没有定义，则新建一个段
    :param doc:
    :param t:
    :return:
    '''
    global last_p
    if last_p is None:
        last_p = new_pa(doc)
    run = new_run(last_p, t)


def add_last_pic(doc, pic, width=None, height=None):
    '''
    在最后段的后面，继续添加图像
    如果最后段没有定义，则新建一个段
    :param doc:
    :param pic:
    :param width:
    :param height:
    :return:
    '''
    global last_p
    if last_p is None:
        last_p = new_pa(doc)
    run = new_run(last_p, '')

    inline_shape = run.add_picture(pic, width, height)
    # 限定图像宽高不能大于15cm。15cm是默认页宽，过大会显示不全
    # 如果超过则等比例缩小，注意缩放的不是图像数据本身
    limit_size = 15
    ori_h_cm = shared.Length(inline_shape.height).cm
    ori_w_cm = shared.Length(inline_shape.width).cm
    if max(ori_h_cm, ori_w_cm) > 15:
        scale = max(ori_h_cm, ori_w_cm) / limit_size
        new_h_cm = ori_h_cm / scale
        new_w_cm = ori_w_cm / scale
        inline_shape.height = shared.Cm(new_h_cm)
        inline_shape.width = shared.Cm(new_w_cm)


def analysis_content(doc, contents: list, ignore_tag_names=[]):
    '''
    分析内容，并建立Doc
    :param doc:
    :param contents:
    :param ignore_tag_names: 用来主动跳过一些标签
    :return:
    '''
    global last_p
    global find_br_count

    for tag in contents:

        # 如果有特定需求，则跳过一些种类的tag
        if tag.name in ignore_tag_names:
            continue

        if tag.name is None:
            t = str(tag)
            # 如果前面找到二个以上的空白行，并且本段文字的长度小于24，则认定本段为章节标题
            # 认错情况略多，还请在输出文档里手动改变
            if find_br_count >= 3 and len(t) < 24:
                add_title(doc, t)
            else:
                add_last_text(doc, t)
            find_br_count = 0

        elif tag.name == 'br':
            # 空白行，目前使用该标记，用来启动新段落
            add_text(doc, '')
            find_br_count += 1

        elif tag.name == 'img':
            # 图像标签
            # 更新方法，如果图像最终下载失败了，那么将会生成一个使用2级标题样式生成链接文本，用来快查
            im_src = tag.attrs['src']
            im_data = get_url(im_src, url)
            if im_data is None:
                # 如果图像下载失败，尝试不使用 ref，因为有的链接使用ref反而无法下载不出来了。。。
                im_data = get_url(im_src)

            if im_data is not None:
                # 图像下载成功
                # doc.add_picture(BytesIO(im_data), width=shared.Cm(15))
                # 因为python-docx的图像解析插件有bug，所以需要手动转换一下图像，用于去除奇怪的问题。
                # python-docx 的图像解析器对元数据的编码处理有bug，可能会导致崩溃
                im = imageio.imread(im_data)

                # python-docx处理png透明图像有bug，百分百报错，这里会将透明图转换成白底不透明图。
                # 注意要手动处理透明通道和白底图像混合，直接截取RGB通道或使用opencv的RGBA转RGB都会导致透明通道处理结果不正确。
                # has_alpha = False
                if im.shape[-1] == 3:
                    im = cv2.cvtColor(im, cv2.COLOR_RGB2BGR)
                elif im.shape[-1] == 4:
                    im = cv2.cvtColor(im, cv2.COLOR_RGBA2BGRA)
                    alpha = im[..., -1:].astype(np.float32) / 255
                    white = np.full([*im.shape[:2], 1], 255, np.float32)
                    im = im.astype(np.float32) * alpha + white * (1 - alpha)
                    im = np.clip(im, 0, 255).astype(np.uint8)

                im_data = cv2.imencode('.jpg', im)[1]

                add_last_pic(doc, BytesIO(im_data))
                analysis_content(doc, tag.contents)
                find_br_count = 0
            else:
                # 如果图像下载失败，则以二级标题方式显示大大的图像链接，方便手动补漏
                add_failure_url(doc, im_src)

        elif tag.name == 'p':
            # 段落标签
            add_text(doc, '')
            analysis_content(doc, tag.contents)
            add_text(doc, '')
            find_br_count = 0

        elif tag.name == 'span':
            # 字体特效字段，但不解析特效
            analysis_content(doc, tag.contents)
            find_br_count = 0

        elif tag.name == 'div':
            # 分隔块标签
            # 如果有隐藏风格，则跳过该段
            if 'style' in tag.attrs:
                if tag.attrs['style'] == 'display: none;':
                    continue

            analysis_content(doc, tag.contents)
            last_p = None
            find_br_count = 0

        elif tag.name == 'a':
            analysis_content(doc, tag.contents)
            find_br_count = 0

        elif tag.name == 'b':
            analysis_content(doc, tag.contents)
            find_br_count = 0

        elif tag.name == 'em':
            analysis_content(doc, tag.contents)
            find_br_count = 0

        else:
            # 对于未知块，直接分析内部内容，同时打印警告。
            print('Warning! Unknow tag', tag)
            analysis_content(doc, tag.contents)
            # raise RuntimeError('Error! Unknow tag.')


# 开始解析
analysis_content(document, content.contents)

# for t in content.contents:
#     print(t)

# 输出文件
document.save(out_path)

print('Success')
